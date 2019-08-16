from docx import Document
from docx.shared import Inches, Pt, RGBColor, Length
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.dml import MSO_THEME_COLOR
from docx.dml.color import ColorFormat
from docx.text.run import Font, Run
from docx.enum.text import WD_ALIGN_PARAGRAPH #does not work for ryans pycharm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT #this one works for ryans pycharm
from docx.oxml.shared import OxmlElement, qn
from copy import deepcopy
import datetime
import textwrap

#pip3 install python-docx
#pip3 install DateTime

# This class takes a template file and then fills in the correct data. It then will
# get emailed for approval.
# template file must be in the same directory or its location specified.
# To instance form gen one should provide the correct OC insturctor name, OC department,
# military_course, and the oc_course

#creating a FileGen object will fill out the top of the form.

#ex:

#fg = FileGen.FileGen(OC_Course, JST_Course, Reviewer)
#fg.find_split_and_copy(len(oc_learning_outcomes))
#for oc, jst in comparison_dict.items():
#       fg.like_outcome_tables(oc, jst)

class FileGen:

        # initializing variables to be used throught the generator
    def __init__(self, oc_course, jst_course, reviewer):# oc_course, jst_course, reviewer
        # test.docx should be the name of the template file.
        self.doc = Document('dbadmin/CourseComparisonMaterial/MCE-TEMPLATE-5.13.2019.docx')# Form_Template_Version_5_13_2019.docx is the other template.
        self.total_tables = self.doc.tables  # locating document tables
        # locating correct table for course comparason
        self.info_table = self.total_tables[0]
        #self.compare_table, self.first_comp_row = self.find_compare_tables() 
        #print(self.first_comp_row)
        self.holder_table = self.total_tables[1]
        self.nested_tables = self.holder_table.cell(0,0).tables
        self.list_holder_tables = [self.holder_table]
        self.compare_tables, self.first_comp_row, self.oc_course_name_cells = self.find_compare_tables()
        self.comp_table = self.compare_tables[0]
        self.table_iterator = 0
        self.update_compare_info()
        #
        self.date_cell = self.info_table.cell(0, 0)
        self.military_course_cell = self.info_table.cell(0, 1)
        # data for filling in top of doc.
        self.instructor_name_cell = self.info_table.cell(1, 0)
        self.oc_course_cell = self.info_table.cell(1, 1)
        self.instr_name = reviewer.name
        self.dep_name = reviewer.department
        self.m_course = jst_course.number
        self.oc_course = oc_course.number
        self.oc_description = oc_course.description
        self.mc_description = jst_course.description
        self.oc_course_name = oc_course.name
        self.mc_course_name = jst_course.name
        # bellow marks out the columns for both military and olivet course
        # outcomes.
        self.mc_column = 1
        self.oc_row = 3
        self.oc_column = 0
        self.total_columns_added = 1 # this should actually be total rows added. will change eventualy.
        self.total_oc_course = 0
        self.total_columns = 6
        # used for sugestive guesses 
        self.no_match = 33
        self.moderate_match = 67
        self.strong_match = 100
        self.line_spacing = 12
        self.remove_last_row = False
        
        
        self.check_table_style()
        self.fill_course_info()

        self.table_iterator = -1
        self.update_compare_info()
        self.ran_find_split_and_copy = False

    # called at the object instance it can be called again if there are more tables created. updating will change what table info is added to.
    def update_compare_info(self):
        self.comp_table = self.compare_tables[self.table_iterator]
        self.comp_rows = self.comp_table.rows
        self.comp_last_row = self.comp_rows[-1]
        self.comp_copy_row = self.comp_rows[1]
        self.mc_outcome_cell = self.comp_table.cell(1,1)
        self.comp_row = self.first_comp_row
        self.tbl = self.comp_table._tbl
        self.mc_row = len(self.comp_rows)-1

    # function used for finding the compare table or tables in a form.
    def find_compare_tables(self):
        tables = []
        cells = []
        f_comp_row = 0
        iterator = 0
        for tbl in self.total_tables:
            if len(tbl.cell(0,0).tables) > 0:# this table holds a nested table
                for table in tbl.cell(0,0).tables:
                    iterator = 1
                    rows = table.rows
                    found_comp_row = False
                    for row in rows:
                        row_cells = row.cells
                        cell_iterator = 0
                        for cell in row_cells:
                            if row_cells[0] == cell and cell_iterator == 1:# looking for the first cell in the nested table
                                    tables.append(table)
                                    cells.append(cell)
                            if cell.paragraphs[0].text == "" and found_comp_row == False:# the first free cell is the first comparison row
                                f_comp_row = iterator
                                found_comp_row = True
                            cell_iterator += 1
                        iterator += 1
        return tables, f_comp_row-1, cells

    # if the table style is different, added rows might not have borders or different borders.
    def check_table_style(self):
        if self.comp_table.style != 'Table Grid':
            self.comp_table.style = 'Table Grid'

    #used to generate more rows for the use of comparason. first adds row to the bottom of the table
    #and then moves it to the correct location.
    def add_row_at(self, location, border=0):
        new_row = self.comp_table.add_row()
        tr = new_row._tr
        self.remove_border_last_row(border=border)# the border is removed based on where it will be incerted into the table.
        if location != "end":# this gives the method the ability to add rows to the end of the table.
            self.tbl.insert(5 + location, tr)#6 meens insert it at the 4th row. 7-5 e.t.c.
            self.total_columns_added += 1
        
        
    # used in emove_order_ast_ow() to find the last row that is created in __Add_Row()
    def getLastRow(self, rows):
        last = rows[-1]
        return last                

    # used to remove the borders of cells, either top or bottom.
    def remove_border(self, row, border=0):
        for cell in row:
            tcPr = cell.tcPr
            tcBorders = OxmlElement('w:tcBorders')
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'nil')
        
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'nil')
        
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'nil')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), 'auto')

            right = OxmlElement('w:right')
            right.set(qn('w:val'), 'nil')

            if border == 1:
                tcBorders.append(top)
            if border == 2:
                tcBorders.append(bottom)
            if border == 0:
                tcBorders.append(top)
                tcBorders.append(bottom)
            tcPr.append(tcBorders)

    # the border is removed before it is incerted into to correct spot in the table.
    def remove_border_last_row(self, border=0):
        self.compare_row = self.comp_rows[self.comp_row]
        rows = self.tbl.getchildren()
        last_row = self.getLastRow(rows)
        self.remove_border(last_row, border=border)

    # used for deleting rows.
    def remove_row(self, row):
        tr = row._tr
        self.tbl.remove(tr)

    # called before saving the program in self.save_doc, this function removes the last un-needed row.
    def remove_empty_row(self):
        empty_row = self.comp_rows[-4]
        self.remove_row(empty_row)
               
    # will add checkboxes the the correct columns
    def add_checkbox(self, jst_outcome, percent):
        column_check_add = self.comp_row

        for i in range(1, len(self.comp_table.columns)):
            cell_check_add = self.compare_row.cells[i]
            para = cell_check_add.paragraphs[0]
            run = para.add_run("\u2610")
            font = run.font
            self.sugested_check(percent, font, i, para, run)
            para.paragraph_format.line_spacing = Pt(self.line_spacing)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER #if WD_ALIGN_PARAGRAPH doesnt work for you switch to WD_PARAGRAPH_ALIGNMENT

    # a percentage determined by nltk to highlight a sugested box to check.
    def sugested_check(self, percent, font, row, para, run):
        if(percent >= 1 and percent <= self.no_match and row == 1):
            #self.highlight(font)
            self.check_mark(run)
        if( percent > self.no_match and percent <= self.moderate_match and row == 2):
            #self.highlight(font)
            self.check_mark(run)
        if(percent > self.moderate_match and percent <= self.strong_match and row == 3):
            #self.highlight(font)
            self.check_mark(run)
        para.paragraph_format.line_spacing = Pt(self.line_spacing)

    # adds a checkmark for sugestive checking.
    def check_mark(self, run):
        run.clear()
        font = run.font
        font.color.theme_color = MSO_THEME_COLOR.TEXT_1
        font.color.rgb = RGBColor(211,211,211)# makes the checkmards light gray
        run.add_text("\u2713")# this is the unicode for checkmark symbol.

    # used for highlighting runs.
    def highlight(self, font):
        font.highlight_color = WD_COLOR_INDEX.GRAY_25

    # called during initualization fills in the top of the table.
    def fill_course_info(self):
        now = datetime.datetime.now()
        self.date_cell.text = "Date of Initiation:\n" + \
            str(now.month) + "-" + str(now.day) + "-" + str(now.year)
        self.military_course_cell.text = "JST or AU course for which Credit/Equivalency is sought:\n" + self.m_course + " - " + self.mc_course_name + "\n" + self.mc_description
        self.instructor_name_cell.text = "Evaluator Name:\n" + \
            self.instr_name + "\nDepartment:\n" + self.dep_name
        self.oc_course_cell.text = "Olivet College course being considered for possible equivalency:\n" + self.oc_course + " - " + self.oc_course_name + "\n" + self.oc_description

    # used for entering individual Olivet course outcomes
    def olivet_course_outcomes(self, c_outcome):
        self.total_oc_course +=1
        self.oc_outcome_cell = self.oc_course_name_cells[self.compare_tables.index(self.comp_table)-1]
        para = self.oc_outcome_cell.paragraphs[0]
        para.add_run("\n" + c_outcome)

    # used for entering individual Military course outcomes.
    # used when the user wants to move to the next cell.
    def jst_outcomes(self, create_row, jst_outcome, percent, new_row=False):
        self.compare_row = self.comp_rows[self.comp_row]
        self.mc_outcome_cell = self.compare_row.cells[0]
        para = self.mc_outcome_cell.paragraphs[0]
        para.add_run(jst_outcome)
        para.paragraph_format.line_spacing = None
        para.paragraph_format.line_spacing = Pt(self.line_spacing)
        if create_row == True:
            if new_row==False:
                self.add_row_at(self.total_columns_added, border=0)
            else:
                self.add_row_at(self.total_columns_added, border=0)
            self.comp_row += 1
        if create_row == False and len(self.compare_tables) == 1:
            self.remove_last_row = True
            self.add_row_at(self.total_columns_added, border=0)
            self.comp_row += 1
        self.add_checkbox(jst_outcome, percent)

    # adds both the Olivet course outcomes with there coresponding military course outcomes.
    # c_outcomes would be just a string for the Olivet college outcome and then jst_outcomes would
    # be the coresponding dictionary that holds percentages.
    def like_outcomes(self, c_outcome, jst_outcome):           
        self.olivet_course_outcomes(c_outcome)
        iterator = 1
        for outcome in jst_outcome:
            if iterator < len(jst_outcome):
                if len(jst_outcome) == iterator+1:
                    self.jst_outcomes(True, outcome[0], outcome[1], new_row=True)
                else:
                    self.jst_outcomes(True, outcome[0], outcome[1])
            else:
                self.jst_outcomes(False, outcome[0], outcome[1], new_row=True)
            iterator +=1

     # this is called after find_split_and_copy() creats the correct tables
    # !!make sure you are useing the correct format when using this method!!
    def like_outcome_tables(self, c_outcome, jst_outcome):
        self.check_find_split_and_copy()# Checking to see if the tables where generated.
        self.update_compare_info()# used to move on to the next table after one is completed.
        self.table_iterator += 1
        self.total_columns_added = "end" # each comparison now has its own table so we down have to incert rows at a point.
        self.like_outcomes(c_outcome, jst_outcome)
        #self.remove_border(self.tbl.getchildren()[2])

        self.comp_row = 1

    # will be used if we decide to implement emailing the form.
    def email_doc(self):
        pass

    def remove_boarder_first_cell(self):
        for table in self.compare_tables:
            tbl = table._tbl
            self.__remove_border(tbl.getchildren()[+3])

    def move_table_after(self, table, paragraph):
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)

    def copy_table_after(self, ctable, atable):
        ctbl, atbl = ctable._tbl, atable._tbl
        new_tbl = deepcopy(ctbl)
        atbl.addnext(new_tbl)
        return new_tbl

    def move_para_after(self, table):
        new_paragraph = self.doc.add_paragraph("")
        p = new_paragraph._p
        table.addnext(p)

    def update_tables(self): # used to update the total_tables property
        self.total_tables = self.doc.tables

    # tables should be total number of oc_outcomes.
    def find_split_and_copy(self, tables): # designed to create as many tables as there are oc outcomes.
        self.ran_find_split_and_copy = True
        iterator = tables-1
        for x in range(0,iterator):# goes through and makes the enough tables to fit each oc_outcome.
            #new_tbl = self.copy_table_after(self.comp_table, self.comp_table)
            new_tbl = self.copy_table_after(self.holder_table, self.holder_table)
            self.move_para_after(new_tbl)
            iterator -= 1
            self.update_tables()
        self.move_para_after(self.holder_table._tbl)
        self.compare_tables, self.first_comp_row, self.oc_course_name_cells = self.find_compare_tables()
        self.update_compare_info()

    # checks to see if tables have been generated.
    def check_find_split_and_copy(self):
        if self.ran_find_split_and_copy == False:
            raise HaveToCreateTables("Have to call find_split_and_copy method first")

    # used to save the document. Must call this to save document.
    def save_doc(self, doc_name='Test-Saved.docx'):
        #self.remove_boarder_first_cell()
        if self.remove_last_row == True:
            self.remove_empty_row()
        self.doc.save(doc_name)

        
# this class is for throwing a error if the like_outcome_table() method is used before the tables are generated using find_split_and_copy()
class HaveToCreateTables(Exception):
    

    def __init__(self, message):
        super().__init__(message)
        self.message = message
        
                

                


                
        